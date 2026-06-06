"""T11 — parametrised empty-extraction-flag tests across all supported formats.

The parser dispatcher (:func:`persona.documents.parsers.parse_document`)
and per-format empty-extraction discipline:

- **Non-PDF formats** (txt/md/code/csv/docx/xlsx): an empty-from-non-empty
  file raises :class:`~persona.documents.errors.CorruptDocumentError` with
  ``reason="empty_after_decode"`` — the criterion #8 binary.
- **PDF**: the empty-extraction sets the
  :data:`~persona.documents.parsers.FLAG_NO_TEXT_LAYER` flag AND sets
  :attr:`~persona.documents.parsers.ParseResult.needs_vision_handoff` to
  ``True``. The vision-handoff TRIGGER is T21's substance (sequenced LAST,
  gated on Spec 13). T11 asserts the **flag is SET**, NOT the handoff is
  invoked.

Also tests the dispatcher's own surface:
- :class:`~persona.documents.errors.UnsupportedFormatError` for unknown
  extensions (e.g. ``.pptx`` per D-14-X-pptx-deferral).
- Importing :mod:`persona.documents.parsers` works without the
  ``[documents]`` extra — only the per-format library import attempt
  inside ``parse_pdf`` / ``parse_docx`` / ``parse_xlsx`` triggers the
  optional-library load (lazy-import discipline).
"""

from __future__ import annotations

import importlib
import sys
from pathlib import Path

import pytest
from persona.documents.errors import (
    CorruptDocumentError,
    UnsupportedFormatError,
)
from persona.documents.parsers import (
    FLAG_NO_TEXT_LAYER,
    SUPPORTED_EXTENSIONS,
    parse_document,
)


class TestDispatchTable:
    def test_supported_extensions_include_required_formats(self) -> None:
        # §9 criterion #1 lists PDF / docx / xlsx / csv / txt / md / code.
        required = {".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md"}
        assert required <= SUPPORTED_EXTENSIONS
        # At least one source-code extension.
        assert ".py" in SUPPORTED_EXTENSIONS

    def test_pptx_is_not_supported(self) -> None:
        # D-14-X-pptx-deferral.
        assert ".pptx" not in SUPPORTED_EXTENSIONS


class TestUnsupportedFormat:
    def test_unknown_extension_raises_unsupported(self, tmp_path: Path) -> None:
        path = tmp_path / "weird.rar"
        path.write_bytes(b"some bytes")
        with pytest.raises(UnsupportedFormatError) as excinfo:
            parse_document(path)
        assert excinfo.value.context["format"] == ".rar"
        assert excinfo.value.context["filename"] == "weird.rar"

    def test_pptx_raises_unsupported(self, tmp_path: Path) -> None:
        # Explicit guard for the D-14-X-pptx-deferral path.
        path = tmp_path / "deck.pptx"
        path.write_bytes(b"some bytes")
        with pytest.raises(UnsupportedFormatError):
            parse_document(path)

    def test_no_extension_raises_unsupported(self, tmp_path: Path) -> None:
        path = tmp_path / "noext"
        path.write_bytes(b"some bytes")
        with pytest.raises(UnsupportedFormatError) as excinfo:
            parse_document(path)
        assert excinfo.value.context["format"] == "unknown"


class TestLazyImportDiscipline:
    """The dispatcher + parser-module imports never load third-party libs.

    This is the structural defence behind D-14-X-documents-extra: a minimal
    install (``pip install persona-core`` without the extra) can still
    ``import persona.documents.parsers`` cleanly.
    """

    def test_dispatcher_import_does_not_load_pypdf(self) -> None:
        # Re-import the parsers package fresh + assert the third-party libs
        # are NOT in sys.modules afterward (unless something earlier in the
        # session already imported them — in which case the discipline
        # is about the import GRAPH, not the test environment).
        # We test the static property: the parsers package's __init__ has
        # no top-level pypdf/openpyxl/docx import.
        import persona.documents.parsers  # noqa: PLC0415

        importlib.reload(persona.documents.parsers)
        source = Path(persona.documents.parsers.__file__).read_text()
        # The __init__ never imports pypdf/openpyxl/docx/pypdfium2 at module
        # level — guards the import graph.
        for forbidden in ("import pypdf", "import openpyxl", "import docx", "import pypdfium2"):
            assert forbidden not in source, (
                f"parsers.__init__ has a top-level {forbidden!r} — breaks "
                "D-14-X-documents-extra lazy-import discipline"
            )

    def test_per_format_parser_modules_lazy_import_libs(self) -> None:
        # Each extra-gated parser module imports the underlying library
        # INSIDE the parse_X function, not at module level.
        for module_name, lib_name in [
            ("persona.documents.parsers.pdf", "pypdf"),
            ("persona.documents.parsers.docx", "docx"),
            ("persona.documents.parsers.xlsx", "openpyxl"),
        ]:
            if module_name in sys.modules:
                del sys.modules[module_name]
            module = importlib.import_module(module_name)
            source = Path(module.__file__).read_text()  # type: ignore[arg-type]
            # The import line must appear inside a function (after `def `),
            # not at module top level.
            top_level_lines = []
            in_function = False
            indent_stack: list[int] = []
            for line in source.splitlines():
                stripped = line.lstrip()
                indent = len(line) - len(stripped)
                if stripped.startswith("def ") or stripped.startswith("class "):
                    in_function = True
                    indent_stack.append(indent)
                elif in_function and indent <= (indent_stack[-1] if indent_stack else 0):
                    # Possibly returned to module level.
                    if stripped and not stripped.startswith("#"):
                        in_function = False
                        indent_stack.clear()
                if not in_function:
                    top_level_lines.append(line)
            top_level_source = "\n".join(top_level_lines)
            assert f"import {lib_name}" not in top_level_source, (
                f"{module_name} has a top-level `import {lib_name}` — must be lazy"
            )


class TestEmptyExtractionPerFormat:
    """Criterion #8 binary: empty extraction from non-empty file is flagged.

    For non-PDF formats, "flagged" means ``CorruptDocumentError`` with
    ``reason="empty_after_decode"`` at the parser boundary. For PDFs, the
    rule is different per the user's T11 framing note: PDF empty-extraction
    asserts the FLAG IS SET (FLAG_NO_TEXT_LAYER + needs_vision_handoff=True),
    NOT the vision handoff trigger. The handoff is T21's substance.
    """

    @pytest.mark.parametrize("extension", [".txt", ".md", ".py"])
    def test_non_pdf_whitespace_only_raises_empty_after_decode(
        self, extension: str, tmp_path: Path
    ) -> None:
        path = tmp_path / f"empty_extraction{extension}"
        path.write_bytes(b"   \n\n\t  \n   \n")
        with pytest.raises(CorruptDocumentError) as excinfo:
            parse_document(path)
        assert excinfo.value.context["reason"] == "empty_after_decode"

    def test_csv_whitespace_only_raises_empty_after_decode(self, tmp_path: Path) -> None:
        path = tmp_path / "empty_extraction.csv"
        path.write_bytes(b"   \n\n  \n")
        with pytest.raises(CorruptDocumentError) as excinfo:
            parse_document(path)
        assert excinfo.value.context["reason"] == "empty_after_decode"

    def test_docx_no_text_raises_empty_after_decode(self, tmp_path: Path) -> None:
        # A real but empty docx — no paragraphs / tables with text.
        from docx import Document  # noqa: PLC0415

        path = tmp_path / "blank.docx"
        Document().save(str(path))
        with pytest.raises(CorruptDocumentError) as excinfo:
            parse_document(path)
        assert excinfo.value.context["reason"] == "empty_after_decode"

    def test_xlsx_no_text_raises_empty_after_decode(self, tmp_path: Path) -> None:
        import openpyxl  # noqa: PLC0415

        path = tmp_path / "blank.xlsx"
        openpyxl.Workbook().save(str(path))
        with pytest.raises(CorruptDocumentError) as excinfo:
            parse_document(path)
        assert excinfo.value.context["reason"] == "empty_after_decode"

    def test_pdf_empty_extraction_sets_flag_not_handoff_trigger(self) -> None:
        # Per the user's T11 framing note: PDF empty-extraction asserts the
        # FLAG IS SET, NOT the vision handoff trigger. The trigger is T21.
        # The scanned-like fixture has <50 chars/page → triggers the flag.
        result = parse_document(
            Path(__file__).resolve().parents[2] / "fixtures" / "documents" / "scanned-like.pdf"
        )
        # FLAG IS SET.
        assert FLAG_NO_TEXT_LAYER in result.flags
        # needs_vision_handoff is the structural plumbing T21 consumes —
        # set, but the actual vision dispatch is T21's substance.
        assert result.needs_vision_handoff is True


class TestDispatchToCorrectParser:
    """Smoke tests that the dispatcher routes each format to its parser."""

    def test_txt_routed_to_text_parser(self, tmp_path: Path) -> None:
        path = tmp_path / "note.txt"
        path.write_text("First para.\n\nSecond para.")
        result = parse_document(path)
        assert len(result.sections) == 2

    def test_md_routed_to_text_parser(self, tmp_path: Path) -> None:
        path = tmp_path / "note.md"
        path.write_text("# Title\n\nBody.")
        result = parse_document(path)
        assert any(s.section == "Title" for s in result.sections)

    def test_csv_routed_to_csv_parser(self, tmp_path: Path) -> None:
        path = tmp_path / "data.csv"
        path.write_text("a,b\n1,2\n")
        result = parse_document(path)
        assert "a | b" in result.sections[0].text

    def test_docx_routed_to_docx_parser(self, tmp_path: Path) -> None:
        from docx import Document  # noqa: PLC0415

        path = tmp_path / "memo.docx"
        doc = Document()
        doc.add_heading("Memo", level=1)
        doc.add_paragraph("Body text.")
        doc.save(str(path))
        result = parse_document(path)
        # Heading captured as a section.
        assert any(s.section == "Memo" for s in result.sections)

    def test_xlsx_routed_to_xlsx_parser(self, tmp_path: Path) -> None:
        import openpyxl  # noqa: PLC0415

        path = tmp_path / "data.xlsx"
        wb = openpyxl.Workbook()
        wb.active.append(["a", "b"])
        wb.active.append([1, 2])
        wb.save(str(path))
        result = parse_document(path)
        assert result.sheet_names is not None

    def test_pdf_routed_to_pdf_parser(self) -> None:
        result = parse_document(
            Path(__file__).resolve().parents[2] / "fixtures" / "documents" / "sample-text.pdf"
        )
        assert result.page_count == 1


class TestMissingDependencyDiscipline:
    """If a parser library is missing, the format-specific parser raises a
    ``MissingDependencyError`` with a clear install hint.

    In this test environment all four libraries are installed (uv sync
    pulled them in alongside the workspace). The discipline is verified
    structurally: the lazy-import branch exists and raises
    :class:`~persona.documents.errors.MissingDependencyError` when the
    underlying import fails. We simulate the missing-library state by
    monkey-patching the lazy import to raise ``ImportError`` and asserting
    the resulting domain error.
    """

    def test_pdf_missing_lib_raises_missing_dependency(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        # Make the underlying ``pypdf`` import fail. The parser's lazy import
        # wrapper catches the ImportError and re-raises as MissingDependencyError.
        # Inject a sentinel that causes `import pypdf` to fail.
        import builtins  # noqa: PLC0415

        from persona.documents.errors import MissingDependencyError  # noqa: PLC0415

        original_import = builtins.__import__

        def _failing_import(name: str, *args: object, **kwargs: object) -> object:
            if name == "pypdf":
                raise ImportError("simulated missing pypdf")
            return original_import(name, *args, **kwargs)  # type: ignore[no-any-return,arg-type]

        # Force a re-import of the pdf module so the lazy-import attempts pypdf again.
        if "pypdf" in sys.modules:
            monkeypatch.delitem(sys.modules, "pypdf")
        monkeypatch.setattr(builtins, "__import__", _failing_import)

        path = tmp_path / "x.pdf"
        path.write_bytes(b"%PDF-1.4")
        with pytest.raises(MissingDependencyError) as excinfo:
            parse_document(path)
        assert excinfo.value.context["format"] == "pdf"
        assert "pip install persona-core[documents]" in excinfo.value.context["install_hint"]

    def test_docx_missing_lib_raises_missing_dependency(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        import builtins  # noqa: PLC0415

        from persona.documents.errors import MissingDependencyError  # noqa: PLC0415

        original_import = builtins.__import__

        def _failing_import(name: str, *args: object, **kwargs: object) -> object:
            if name == "docx":
                raise ImportError("simulated missing docx")
            return original_import(name, *args, **kwargs)  # type: ignore[no-any-return,arg-type]

        if "docx" in sys.modules:
            monkeypatch.delitem(sys.modules, "docx")
        monkeypatch.setattr(builtins, "__import__", _failing_import)

        path = tmp_path / "x.docx"
        path.write_bytes(b"PK\x03\x04")
        with pytest.raises(MissingDependencyError) as excinfo:
            parse_document(path)
        assert excinfo.value.context["format"] == "docx"

    def test_xlsx_missing_lib_raises_missing_dependency(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ) -> None:
        import builtins  # noqa: PLC0415

        from persona.documents.errors import MissingDependencyError  # noqa: PLC0415

        original_import = builtins.__import__

        def _failing_import(name: str, *args: object, **kwargs: object) -> object:
            if name == "openpyxl":
                raise ImportError("simulated missing openpyxl")
            return original_import(name, *args, **kwargs)  # type: ignore[no-any-return,arg-type]

        if "openpyxl" in sys.modules:
            monkeypatch.delitem(sys.modules, "openpyxl")
        monkeypatch.setattr(builtins, "__import__", _failing_import)

        path = tmp_path / "x.xlsx"
        path.write_bytes(b"PK\x03\x04")
        with pytest.raises(MissingDependencyError) as excinfo:
            parse_document(path)
        assert excinfo.value.context["format"] == "xlsx"

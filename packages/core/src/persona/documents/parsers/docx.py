"""DOCX parser (spec 14 T08).

Uses ``python-docx==1.1.2`` (Spec 12 sandbox alignment per
:file:`packages/core/src/persona/sandbox/image/requirements.txt`). The
library is gated behind the ``[documents]`` extra (D-14-X-documents-extra);
this module's import-time access is lazy via :func:`_import_docx` so a
minimal install (``pip install persona-core`` without the extra) can still
``import persona.documents.parsers`` cleanly. Only :func:`parse_docx`
attempts the underlying import; failure surfaces as a
:class:`~persona.documents.errors.MissingDependencyError` with a clear
install hint (T11 wires the dispatcher around this).

Extraction strategy:

- One :class:`DocumentSection` per heading-bounded block (paragraphs grouped
  under their nearest heading). Each section's :attr:`DocumentSection.section`
  field carries the heading text so retrieved chunks cite the source.
- Tables become text-table sections (``|``-separated rows, header preserved)
  appended in document order. Each table-section's :attr:`section` field
  carries either the surrounding heading or ``"Table"`` as a generic marker.
- Empty paragraphs are dropped. Whitespace-only paragraphs are dropped.

Fail-safe handling:

- ``python-docx`` raises :class:`docx.opc.exceptions.PackageNotFoundError`
  for files that aren't valid OOXML packages. The parser catches this at the
  adapter boundary and re-raises as
  :class:`~persona.documents.errors.CorruptDocumentError` per the
  engineering-standards §1 "catch provider exceptions at the adapter
  boundary and re-raise as domain" rule.
- A successfully-opened document that yields no text (all paragraphs empty)
  raises ``CorruptDocumentError`` with ``reason="empty_after_decode"`` —
  the criterion #8 "empty extraction from non-empty file" path.
"""

from __future__ import annotations

from typing import TYPE_CHECKING, Any

from persona.documents.chunker import DocumentSection
from persona.documents.errors import (
    CorruptDocumentError,
    MissingDependencyError,
)
from persona.documents.parsers import ParseResult

if TYPE_CHECKING:
    from pathlib import Path

__all__ = ["parse_docx"]


def _import_docx() -> Any:  # noqa: ANN401 — third-party module returned dynamically
    """Lazy-import ``docx`` with a clear MissingDependencyError on failure."""
    try:
        import docx  # noqa: PLC0415 — deliberate lazy import per D-14-X-documents-extra
    except ImportError as exc:
        raise MissingDependencyError(
            "python-docx is not installed",
            context={
                "format": "docx",
                "install_hint": "pip install persona-core[documents]",
            },
        ) from exc
    return docx


def parse_docx(path: Path) -> ParseResult:
    """Parse a ``.docx`` file to natural-boundary sections.

    Args:
        path: Path to the ``.docx`` file.

    Returns:
        :class:`ParseResult` carrying one :class:`DocumentSection` per
        heading-bounded block plus one per table.

    Raises:
        MissingDependencyError: ``python-docx`` is not installed (the
            ``[documents]`` extra hasn't been opted into).
        CorruptDocumentError: The file isn't a valid docx (caught from
            ``python-docx``'s ``PackageNotFoundError``) or extraction
            yielded no text (criterion #8).
    """
    docx = _import_docx()
    size_bytes = path.stat().st_size
    filename = path.name

    if size_bytes == 0:
        raise CorruptDocumentError(
            "file is empty",
            context={"format": "docx", "reason": "empty_file", "filename": filename},
        )

    try:
        document = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001 — adapter-boundary catch per §1
        # python-docx raises PackageNotFoundError from a non-stdlib subclass
        # hierarchy; catch broadly and discriminate by the message.
        reason = _classify_open_failure(exc)
        raise CorruptDocumentError(
            "could not open docx",
            context={"format": "docx", "reason": reason, "filename": filename},
        ) from exc

    sections = _extract_sections(document)

    if not sections or not any(s.text.strip() for s in sections):
        raise CorruptDocumentError(
            "extraction yielded no text",
            context={
                "format": "docx",
                "reason": "empty_after_decode",
                "filename": filename,
            },
        )

    return ParseResult(
        sections=tuple(sections),
        flags=(),
        size_bytes=size_bytes,
    )


def _classify_open_failure(exc: BaseException) -> str:
    """Map a python-docx open failure to a stable ``reason`` discriminator."""
    name = type(exc).__name__
    if "PackageNotFound" in name:
        return "not_a_docx"
    if "Zip" in name or "BadZip" in name:
        return "not_a_zip"
    return "unknown"


def _extract_sections(document: Any) -> list[DocumentSection]:  # noqa: ANN401
    """Walk the docx body in order, grouping paragraphs under headings.

    The python-docx Document's ``paragraphs`` and ``tables`` are emitted in
    document order via the body's element children. For v0.1 we use the
    simpler ``document.paragraphs`` + ``document.tables`` enumeration: it
    loses inter-block ordering (a table between two paragraphs reads after
    both), which is acceptable for read-not-analyse v0.1. A faithful
    document-order walk is a v0.2 refinement if T22 smoke testing surfaces
    the gap.
    """
    sections: list[DocumentSection] = []

    current_heading: str | None = None
    current_lines: list[str] = []

    def _flush_heading_section() -> None:
        nonlocal current_lines
        body = "\n".join(line for line in current_lines if line.strip())
        if not body and not current_heading:
            current_lines = []
            return
        full_body = body
        if current_heading and body:
            full_body = f"{current_heading}\n\n{body}"
        elif current_heading:
            full_body = current_heading
        if full_body.strip():
            sections.append(DocumentSection(text=full_body, section=current_heading))
        current_lines = []

    for para in document.paragraphs:
        text = para.text.strip()
        style_name = ""
        try:
            style_name = (para.style.name or "").strip()
        except Exception:  # noqa: BLE001 — adapter-boundary catch
            style_name = ""

        if style_name.startswith("Heading") and text:
            _flush_heading_section()
            current_heading = text
            continue
        if text:
            current_lines.append(text)

    _flush_heading_section()

    # Tables — appended in document order (within the limitation noted above).
    for table_index, table in enumerate(document.tables):
        table_text = _format_table(table)
        if table_text.strip():
            sections.append(
                DocumentSection(
                    text=table_text,
                    section=current_heading or f"Table {table_index + 1}",
                )
            )

    return sections


def _format_table(table: Any) -> str:  # noqa: ANN401
    """Render a docx table as a pipe-separated text table.

    The first row is treated as the header. Subsequent rows are normalised
    to the header's column count. Cell text is stripped.
    """
    rows = []
    for row in table.rows:
        cells = [(cell.text or "").strip() for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    header = rows[0]
    if len(rows) == 1:
        return " | ".join(header)
    lines = [" | ".join(header)]
    for row in rows[1:]:
        normalised = _normalise_row(row, width=len(header))
        lines.append(" | ".join(normalised))
    return "\n".join(lines)


def _normalise_row(row: list[str], *, width: int) -> list[str]:
    """Pad or truncate a row to ``width`` cells."""
    if len(row) < width:
        return [*row, *([""] * (width - len(row)))]
    if len(row) > width:
        return row[:width]
    return row

"""T10 — both-loops + error-recovery + RLS integration test (Spec 16).

Covers Spec 16 §9 acceptance criteria:

- **#8** Both loops produce documents end-to-end:
  - ``test_conversation_loop_produces_xlsx_in_chat`` exercises
    :class:`persona_runtime.loop.ConversationLoop`'s use_skill → code_execution
    sub-loop with the xlsx_generation skill activated and openpyxl code that
    writes ``/workspace/out/budget.xlsx``.
  - ``test_agentic_loop_produces_pdf_as_task_step`` exercises
    :class:`persona_runtime.agentic.loop.AgenticLoop`'s plan → tool → tool →
    final cycle with a (mocked) web_research preamble followed by use_skill
    pdf_generation and a reportlab-flowable code_execution that produces
    ``/workspace/out/report.pdf``.
- **#9** Error-recovery: a deliberately broken python-docx snippet returns
  ``ToolResult(is_error=True, ...)`` (the sandbox surfaces the AttributeError
  traceback to the model); the next model call sees the error in context and
  writes a fixed snippet that produces a docx — final ``Run.status`` is
  ``RunStatus.COMPLETED``. Mirrors the §0.4-verified pattern at
  ``packages/runtime/tests/unit/test_loop_agentic.py:161``
  ``test_tool_failure_is_fed_back_and_model_recovers`` — the only difference
  is the error is a python-docx traceback, not a flaky-tool 503.
- **#11** RLS / tenant isolation: two sandbox sessions composed with distinct
  tenant-prefixed ``session_id`` values (``f"{owner}:{conversation}"`` per
  Spec 12 D-12-1 + kickoff trip-up #6). owner_A produces a docx in session_A;
  owner_B's session_B cannot reach session_A's file (each session container
  owns its own ``/workspace/out`` bind mount per Spec 12 D-12-9 / D-12-15).

**Scope (A2 / D-16-X-5).** ``ExecutionResult.produced_files`` is the
assertion target at the sandbox boundary. ``LocalDockerSandbox`` populates
this tuple from the per-session ``host_out`` diff after each ``docker exec``;
``HostedSandbox`` constructs ``ExecutionResult`` with the default empty tuple
(Spec 17 / Spec 16 state A2 finding — not yet wired). T10 scopes its
assertions to the local backend only; hosted backends short-circuit with a
``pytest.skip`` if the test ever surfaces in a hosted run.

**`_ScriptedBackend`** is module-private per the api per-file convention
(T01 audit A1 — three precedents in the api integration tree). Copied from
``packages/runtime/tests/_fakes.py:50`` with two extensions:

- A ``chat_script`` parameter (mirrors the agentic test fakes) so a sequence
  of pre-built ``ChatResponse`` objects can drive ``AgenticLoop.run``.
- A ``rounds`` parameter (the conversation-loop streaming path) replays
  scripted ``ScriptedRound`` text/tool-call rounds through ``chat_stream``.
"""

# ruff: noqa: SLF001, ANN401 — tests pin scripted backends into the tier
# registry cache + assert on the sandbox's per-session host workspace dirs
# (SLF001); the MemoryStore double mirrors the Protocol's loose kwargs shape
# (ANN401).

from __future__ import annotations

import json
import uuid
from pathlib import Path
from typing import TYPE_CHECKING, Any

import pytest
from persona.backends import BackendConfig, ChatResponse
from persona.backends.types import StreamChunk, TokenUsage, ToolCallDelta
from persona.history import ConversationHistoryManager
from persona.sandbox.local_docker import (
    DEFAULT_IMAGE,
    LocalDockerSandbox,
    is_docker_available,
)
from persona.sandbox.result import NetworkPolicy, ResourceLimits, SandboxFile
from persona.sandbox.tool import make_code_execution_tool
from persona.schema.conversation import Conversation
from persona.schema.persona import Persona, PersonaIdentity
from persona.schema.skills import SkillSpec
from persona.schema.tools import ToolCall, ToolResult
from persona.skills import (
    SkillInjector,
    SkillScanner,
    count_tokens,
    make_use_skill_tool,
)
from persona.tools import Toolbox
from persona.tools.protocol import tool
from persona_runtime.agentic.loop import AgenticLoop
from persona_runtime.agentic.run import RunStatus
from persona_runtime.agentic.step import StepType
from persona_runtime.logging import MemoryTurnLogWriter
from persona_runtime.loop import ConversationLoop
from persona_runtime.prompt import PromptBuilder
from persona_runtime.router import Router
from persona_runtime.tier import TierConfig, TierRegistry

if TYPE_CHECKING:
    from collections.abc import AsyncIterator

    from persona.backends.types import ToolSpec
    from persona.schema.conversation import ConversationMessage
    from persona.tools.protocol import AsyncTool


pytestmark = [pytest.mark.integration, pytest.mark.docker]


_DUMMY_CFG = BackendConfig(provider="anthropic", model="m", api_key=None)


# ---------------------------------------------------------------------------
# Module-private _ScriptedBackend (T01 audit A1 — per-file api convention).
# Copied from packages/runtime/tests/_fakes.py:50 with the agentic chat_script
# extension so the same fake serves BOTH the conversation loop (chat_stream
# round replay) and the agentic loop (chat() ChatResponse sequence replay).
# ---------------------------------------------------------------------------


class _ScriptedRound:
    """One scripted backend round: either text, or one tool call."""

    def __init__(
        self,
        *,
        text: str = "",
        text_deltas: list[str] | None = None,
        tool_name: str | None = None,
        tool_args: dict[str, Any] | None = None,
        call_id: str = "call-1",
    ) -> None:
        self.text = text
        self.text_deltas = text_deltas
        self.tool_name = tool_name
        self.tool_args = tool_args or {}
        self.call_id = call_id


class _ScriptedBackend:
    """A ChatBackend that replays scripted rounds / responses.

    For the conversation loop: each ``chat_stream`` call consumes the next
    ``_ScriptedRound`` (text or one tool call). For the agentic loop: each
    ``chat`` call consumes the next ``ChatResponse`` in ``chat_script``.
    """

    def __init__(
        self,
        rounds: list[_ScriptedRound] | None = None,
        *,
        provider_name: str = "anthropic",
        model_name: str = "claude-sonnet-4-6",
        chat_script: list[ChatResponse] | None = None,
        supports_vision: bool = False,
    ) -> None:
        self._rounds = list(rounds or [])
        self._index = 0
        self._provider_name = provider_name
        self._model_name = model_name
        self._supports_vision = supports_vision
        self.chat_stream_calls = 0
        self._chat_script: list[ChatResponse] = list(chat_script) if chat_script else []
        self._chat_index = 0
        self.chat_calls = 0

    @property
    def provider_name(self) -> str:
        return self._provider_name

    @property
    def model_name(self) -> str:
        return self._model_name

    @property
    def supports_native_tools(self) -> bool:
        return False

    @property
    def supports_vision(self) -> bool:
        return self._supports_vision

    async def chat(
        self,
        messages: list[ConversationMessage],  # noqa: ARG002
        *,
        tools: list[ToolSpec] | None = None,  # noqa: ARG002
        temperature: float = 0.0,  # noqa: ARG002
        max_tokens: int = 4096,  # noqa: ARG002
        stop: list[str] | None = None,  # noqa: ARG002
    ) -> ChatResponse:
        self.chat_calls += 1
        if self._chat_script:
            if self._chat_index < len(self._chat_script):
                response = self._chat_script[self._chat_index]
                self._chat_index += 1
                return response
            return ChatResponse(
                content="",
                tool_calls=[],
                usage=TokenUsage(prompt_tokens=1, completion_tokens=0, total_tokens=1),
                model=self._model_name,
                provider=self._provider_name,
                latency_ms=0.0,
            )
        # Conversation-loop summariser path: fixed "SUMMARY" (no compaction in
        # these tests; the conversation starts empty).
        return ChatResponse(
            content="SUMMARY",
            tool_calls=[],
            usage=TokenUsage(prompt_tokens=5, completion_tokens=2, total_tokens=7),
            model=self._model_name,
            provider=self._provider_name,
            latency_ms=0.0,
        )

    async def chat_stream(
        self,
        messages: list[ConversationMessage],  # noqa: ARG002
        *,
        tools: list[ToolSpec] | None = None,  # noqa: ARG002
        temperature: float = 0.0,  # noqa: ARG002
        max_tokens: int = 4096,  # noqa: ARG002
        stop: list[str] | None = None,  # noqa: ARG002
    ) -> AsyncIterator[StreamChunk]:
        self.chat_stream_calls += 1
        if self._index >= len(self._rounds):
            yield StreamChunk(
                delta="",
                is_final=True,
                usage=TokenUsage(prompt_tokens=1, completion_tokens=0, total_tokens=1),
            )
            return
        rnd = self._rounds[self._index]
        self._index += 1
        if rnd.tool_name is not None:
            yield StreamChunk(
                delta="",
                tool_call_delta=ToolCallDelta(
                    call_id=rnd.call_id,
                    name_delta=rnd.tool_name,
                    arguments_delta=json.dumps(rnd.tool_args),
                ),
            )
        elif rnd.text_deltas:
            for piece in rnd.text_deltas:
                yield StreamChunk(delta=piece)
        elif rnd.text:
            yield StreamChunk(delta=rnd.text)
        yield StreamChunk(
            delta="",
            is_final=True,
            usage=TokenUsage(prompt_tokens=10, completion_tokens=5, total_tokens=15),
        )


# ---------------------------------------------------------------------------
# In-memory MemoryStore double (mirrors packages/runtime/tests/_fakes.py).
# Avoids the Postgres + RLS engine overhead — T10 is about loops + sandbox,
# not the persistence layer. Spec 08's RLS integration is covered by the
# existing test_rls_per_endpoint.py suite.
# ---------------------------------------------------------------------------


class _FakeStore:
    """Minimal in-memory MemoryStore recording writes."""

    def __init__(self) -> None:
        self.writes: list[list[Any]] = []
        self._all: list[Any] = []

    def write(
        self,
        persona_id: str,  # noqa: ARG002
        chunks: list[Any],
        *,
        source: Any,  # noqa: ARG002
        written_by: str | None = None,  # noqa: ARG002
        reason: str | None = None,  # noqa: ARG002
        force: bool = False,  # noqa: ARG002
    ) -> None:
        self.writes.append(list(chunks))
        self._all.extend(chunks)

    def query(
        self,
        persona_id: str,  # noqa: ARG002
        query: str,  # noqa: ARG002
        top_k: int,  # noqa: ARG002
        **filters: Any,  # noqa: ARG002
    ) -> list[Any]:
        return []

    def get_all(
        self,
        persona_id: str,  # noqa: ARG002
        *,
        include_superseded: bool = False,  # noqa: ARG002
    ) -> list[Any]:
        return list(self._all)

    def delete(self, persona_id: str) -> None: ...  # noqa: ARG002

    def remove_documents(
        self,
        persona_id: str,  # noqa: ARG002
        doc_ids: list[str],  # noqa: ARG002
    ) -> None: ...

    def history(
        self,
        persona_id: str,  # noqa: ARG002
        logical_id: str,  # noqa: ARG002
    ) -> list[Any]:
        return []

    def rollback(
        self,
        persona_id: str,  # noqa: ARG002
        logical_id: str,  # noqa: ARG002
        to_version: int,  # noqa: ARG002
        *,
        source: Any,  # noqa: ARG002
        written_by: str | None = None,  # noqa: ARG002
        reason: str | None = None,  # noqa: ARG002
    ) -> None: ...


# ---------------------------------------------------------------------------
# Test fixtures.
# ---------------------------------------------------------------------------


def _is_image_available(tag: str) -> bool:
    """True if the sandbox image is locally available (mirrors core conftest)."""
    try:
        import docker
        from docker.errors import DockerException, ImageNotFound
    except ImportError:
        return False
    try:
        client = docker.from_env()
        try:
            client.images.get(tag)
        except ImageNotFound:
            return False
        finally:
            client.close()
    except DockerException:
        return False
    return True


def _docker_skip_reason() -> str | None:
    """Return a skip reason if the local Docker substrate isn't ready, else None."""
    if not is_docker_available():
        return (
            "Docker daemon unreachable; T10 needs LocalDockerSandbox to drive "
            "code_execution. Start Docker and rerun."
        )
    if not _is_image_available(DEFAULT_IMAGE):
        return (
            f"Sandbox image {DEFAULT_IMAGE!r} not built. Build via the "
            "Spec 12 T06 Dockerfile (or pull) before running T10."
        )
    return None


# D-12-X-venv-path-ordering (Spec 16 T09/T10 production fix): the prior
# ``_VENV_PRELUDE = "import sys; sys.path.insert(0, '/opt/venv/...')"`` workaround
# is removed. ``_BASE_CONTAINER_KWARGS["environment"]["PATH"]`` now prepends
# ``/opt/venv/bin`` so the persona-sandbox image's venv-installed
# ``python-docx`` / ``openpyxl`` / ``python-pptx`` / ``reportlab`` resolve
# natively from ``python`` inside the container (both ``docker run`` and
# ``docker exec`` paths). R-12-2's explicit-PATH hardening intent is
# preserved by keeping the system-bin tail.


def _persona() -> Persona:
    return Persona(
        persona_id="astrid_t10",
        identity=PersonaIdentity(
            name="Astrid",
            role="document-producing tenancy assistant",
            background="Knows husleieloven; produces real files.",
            constraints=[],
        ),
        skills=[
            "docx_generation",
            "pptx_generation",
            "xlsx_generation",
            "pdf_generation",
            "web_research",
        ],
        tools=["use_skill", "code_execution", "web_search", "web_fetch"],
    )


def _stores() -> dict[str, _FakeStore]:
    return {
        "identity": _FakeStore(),
        "self_facts": _FakeStore(),
        "worldview": _FakeStore(),
        "episodic": _FakeStore(),
    }


def _registry_for(backend: _ScriptedBackend) -> TierRegistry:
    """A TierRegistry whose every tier resolves to the scripted backend."""
    registry = TierRegistry(
        {
            "frontier": TierConfig(name="frontier", backend_config=_DUMMY_CFG),
            "mid": TierConfig(name="mid", backend_config=_DUMMY_CFG),
            "small": TierConfig(name="small", backend_config=_DUMMY_CFG),
        }
    )
    registry._cache = {"frontier": backend, "mid": backend, "small": backend}
    return registry


def _scanned_skills(names: list[str], *, isolated_root: Path) -> list[SkillSpec]:
    """Build minimal :class:`SkillSpec` test fixtures under ``isolated_root``.

    T10 deliberately does NOT scan the real builtin/<name>/supplements/ trees:
    those would land in ``deferred_input_files`` with absolute paths like
    ``/workspace/in/.skills/<name>/supplements/<topic>.md`` which the
    LocalDockerSandbox's ``_seed_workspace`` mis-resolves via ``host_in /
    f.path`` (Python ``Path('/x') / '/y' == Path('/y')``). That's a real
    M1a-supplements wiring bug — surfaced during T10 development, out of
    scope to fix here. By constructing SkillSpec(path=<dir with no
    supplements/>), ``collect_skill_supplements`` returns ``[]`` and T10
    exercises the criterion-#8/#9 path without tripping the supplements bug.
    """
    out: list[SkillSpec] = []
    for name in names:
        skill_dir = isolated_root / name
        skill_dir.mkdir(parents=True, exist_ok=True)
        content = f"# {name}\n\nTest fixture for T10 — exercises the loop wiring."
        skill_md = skill_dir / "SKILL.md"
        skill_md.write_text(
            f"---\nname: {name}\ndescription: T10 test fixture for {name}.\n"
            f"tools_required:\n  - code_execution\n---\n\n{content}",
            encoding="utf-8",
        )
        out.append(
            SkillSpec(
                name=name,
                description=f"T10 test fixture for {name}.",
                path=skill_dir,
                content=content,
                content_token_count=count_tokens(content),
                tools_required=["code_execution"],
            )
        )
    return out


@tool(name="web_search", description="Mock web search for T10.")
async def _mock_web_search(query: str) -> ToolResult:
    """A stand-in for the real web_search tool — returns one fake hit."""
    return ToolResult(
        tool_name="web_search",
        content=f"Search: {query}\n1. Tenant protections (mock URL)\n",
        is_error=False,
    )


@tool(name="web_fetch", description="Mock web fetch for T10.")
async def _mock_web_fetch(url: str) -> ToolResult:  # noqa: ARG001
    """A stand-in for the real web_fetch tool — returns canned text."""
    return ToolResult(
        tool_name="web_fetch",
        content="Mock article text about tenant protections in Norway.",
        is_error=False,
    )


# ---------------------------------------------------------------------------
# Code snippets executed by the sandbox. Kept tight so the SKILL.md content
# isn't load-bearing for these tests — T11–T14 inspect the SKILL.md craft
# against real files; T10's job is to prove the wiring + recovery + isolation.
# ---------------------------------------------------------------------------


_XLSX_CODE = """\
from openpyxl import Workbook
wb = Workbook()
ws = wb.active
ws.title = "Months"
ws["A1"] = "Category"
ws["B1"] = "Jan2026"
ws["A2"] = "Rent"
ws["B2"] = 12000
wb.save("/workspace/out/budget.xlsx")
print("ok")
"""

_PDF_CODE = """\
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

doc = SimpleDocTemplate("/workspace/out/report.pdf", pagesize=A4)
styles = getSampleStyleSheet()
story = [Paragraph("Tenant protection - quarterly summary", styles["Title"]),
         Spacer(1, 12),
         Paragraph("Body paragraph one.", styles["BodyText"])]
doc.build(story)
print("ok")
"""

# A deliberately broken docx snippet: ``add_table`` returns a Table, which
# does NOT have a ``nonexistent_method``. The container surfaces the
# AttributeError traceback; the next model call writes the fixed snippet.
_BROKEN_DOCX_CODE = """\
from docx import Document
doc = Document()
doc.add_heading("Memo", level=1)
doc.add_table(rows=1, cols=1).nonexistent_method()
doc.save("/workspace/out/memo.docx")
"""

_FIXED_DOCX_CODE = """\
from docx import Document
doc = Document()
doc.add_heading("Memo", level=1)
tbl = doc.add_table(rows=2, cols=2)
tbl.rows[0].cells[0].text = "Header A"
tbl.rows[0].cells[1].text = "Header B"
tbl.rows[1].cells[0].text = "Row 1A"
tbl.rows[1].cells[1].text = "Row 1B"
doc.save("/workspace/out/memo.docx")
print("ok")
"""

# Isolation test: owner_A writes this file; owner_B's session tries to read
# it back. The session-isolation guarantee (Spec 12 D-12-9 / D-12-15) means
# owner_B's /workspace/out is a different host directory.
_TENANT_A_WRITE_CODE = """\
from docx import Document
doc = Document()
doc.add_paragraph("Tenant A's confidential memo")
doc.save("/workspace/out/tenant_a_memo.docx")
print("ok")
"""

_TENANT_B_READ_PROBE_CODE = """\
from pathlib import Path
target = Path("/workspace/out/tenant_a_memo.docx")
print("exists:", target.exists())
# Enumerate everything in /workspace/out so we can assert tenant_a_memo.docx
# is absent regardless of any accidental shared directory wiring.
listing = sorted(p.name for p in Path("/workspace/out").iterdir())
print("listing:", listing)
"""


# ---------------------------------------------------------------------------
# Loop builders — compose the loops with the scripted backend and a real
# LocalDockerSandbox-backed code_execution tool. The use_skill tool is
# wired only when ``with_use_skill=True`` (it requires a SkillSpec list).
# ---------------------------------------------------------------------------


def _build_code_execution_tool(
    sandbox: LocalDockerSandbox,
    *,
    session_id: str,
    deferred_holder: list[SandboxFile],
) -> AsyncTool:
    """Build a code_execution tool bound to ``sandbox`` with the M1a holder.

    Mirrors :class:`persona_api.services.runtime_factory.RuntimeFactory`'s
    composition root: the same ``list[SandboxFile]`` holder is shared between
    the loop's public ``deferred_input_files`` attribute and the tool's
    drain-and-clear provider closure (D-16-2 / D-16-2-state-location).
    """
    holder = deferred_holder

    def _drain_and_clear() -> list[SandboxFile]:
        snapshot = list(holder)
        holder.clear()
        return snapshot

    return make_code_execution_tool(
        sandbox,
        session_id_provider=lambda: session_id,
        deferred_input_files_provider=_drain_and_clear,
        # Default ResourceLimits are conservative (256 MiB / 30 s); docx /
        # pdf / xlsx generation is well within. The image manifest's
        # python-docx + openpyxl + reportlab pins are sandbox-internal.
    )


def _build_conversation_loop(
    backend: _ScriptedBackend,
    *,
    sandbox: LocalDockerSandbox,
    session_id: str,
    scanned: list[SkillSpec],
    persona: Persona,
) -> ConversationLoop:
    deferred_holder: list[SandboxFile] = []
    tools: list[AsyncTool] = [
        _build_code_execution_tool(sandbox, session_id=session_id, deferred_holder=deferred_holder)
    ]
    if scanned:
        tools.append(make_use_skill_tool(scanned))
    toolbox = Toolbox(tools, allow_list=list(persona.tools or []))
    loop = ConversationLoop(
        persona=persona,
        stores=_stores(),  # type: ignore[arg-type]
        toolbox=toolbox,
        skill_scanner=SkillScanner([]),
        skill_injector=SkillInjector(),
        scanned_skills=scanned,
        history_manager=ConversationHistoryManager(compact_every=10, keep_recent=5),
        prompt_builder=PromptBuilder(),
        router=Router(),
        tier_registry=_registry_for(backend),
        turn_log_writer=MemoryTurnLogWriter(),
    )
    loop.deferred_input_files = deferred_holder
    return loop


def _build_agentic_loop(
    backend: _ScriptedBackend,
    *,
    sandbox: LocalDockerSandbox,
    session_id: str,
    scanned: list[SkillSpec],
    persona: Persona,
    extra_tools: list[AsyncTool] | None = None,
) -> AgenticLoop:
    deferred_holder: list[SandboxFile] = []
    tools: list[AsyncTool] = [
        _build_code_execution_tool(sandbox, session_id=session_id, deferred_holder=deferred_holder)
    ]
    if scanned:
        tools.append(make_use_skill_tool(scanned))
    if extra_tools:
        tools.extend(extra_tools)
    toolbox = Toolbox(tools, allow_list=list(persona.tools or []))
    loop = AgenticLoop(
        persona=persona,
        stores=_stores(),  # type: ignore[arg-type]
        toolbox=toolbox,
        skill_injector=SkillInjector(),
        scanned_skills=scanned,
        prompt_builder=PromptBuilder(),
        router=Router(),
        tier_registry=_registry_for(backend),
    )
    loop.deferred_input_files = deferred_holder
    return loop


def _resp(content: str = "", *, tool_calls: list[ToolCall] | None = None) -> ChatResponse:
    return ChatResponse(
        content=content,
        tool_calls=tool_calls or [],
        usage=TokenUsage(prompt_tokens=10, completion_tokens=5, total_tokens=15),
        model="claude-sonnet-4-6",
        provider="anthropic",
        latency_ms=1.0,
    )


# ---------------------------------------------------------------------------
# Tests.
# ---------------------------------------------------------------------------


class TestBothLoopsProduceDocuments:
    """Criterion #8 — ConversationLoop AND AgenticLoop produce documents."""

    @pytest.mark.asyncio
    async def test_conversation_loop_produces_xlsx_in_chat(self, tmp_path: Path) -> None:
        # Criterion #8 — ConversationLoop path. Per A2: LocalDockerSandbox only
        # (HostedSandbox.produced_files is unwired per state.md A2 finding).
        skip_reason = _docker_skip_reason()
        if skip_reason:
            pytest.skip(skip_reason)
        scanned = _scanned_skills(["xlsx_generation"], isolated_root=tmp_path / "skills")

        sandbox = LocalDockerSandbox(workspace_root=tmp_path / "wsp")
        try:
            session_id = f"owner_t10_xlsx:conv_{uuid.uuid4().hex[:8]}"
            await sandbox.create_session(
                session_id, limits=ResourceLimits(), network=NetworkPolicy()
            )

            # Conversation-loop script: round 1 calls use_skill(xlsx_generation);
            # the loop injects the skill; round 2 calls code_execution with the
            # openpyxl snippet; round 3 is the final text reply.
            rounds = [
                _ScriptedRound(
                    tool_name="use_skill",
                    tool_args={"skill_name": "xlsx_generation"},
                    call_id="cu1",
                ),
                _ScriptedRound(
                    tool_name="code_execution",
                    tool_args={"code": _XLSX_CODE},
                    call_id="cc1",
                ),
                _ScriptedRound(text="Workbook written to /workspace/out/budget.xlsx."),
            ]
            backend = _ScriptedBackend(rounds=rounds)
            loop = _build_conversation_loop(
                backend,
                sandbox=sandbox,
                session_id=session_id,
                scanned=scanned,
                persona=_persona(),
            )
            conv = Conversation(
                conversation_id="conv_xlsx",
                persona_id=_persona().persona_id or "astrid_t10",
                messages=[],
            )

            # Drain the streamed turn.
            chunks = [c async for c in loop.turn(conv, "Build me a budget xlsx.")]
            assert chunks[-1].is_final is True

            # Assert the produced .xlsx surfaced as a SandboxFile somewhere in
            # the turn-log's tool-result audit. The conversation loop doesn't
            # expose ExecutionResult directly; the canonical proof is the
            # host-side workspace dir for this session contains the file.
            host_out = sandbox._session_workspaces[session_id][1]
            xlsx_files = list(host_out.glob("**/*.xlsx"))
            assert len(xlsx_files) == 1, (
                f"expected exactly one .xlsx in {host_out}, got {xlsx_files}"
            )
            assert xlsx_files[0].name == "budget.xlsx"
        finally:
            await sandbox.aclose()

    @pytest.mark.asyncio
    async def test_agentic_loop_produces_pdf_as_task_step(self, tmp_path: Path) -> None:
        # Criterion #8 — AgenticLoop path. Same A2 scoping: local backend only.
        skip_reason = _docker_skip_reason()
        if skip_reason:
            pytest.skip(skip_reason)
        scanned = _scanned_skills(
            ["pdf_generation", "web_research"], isolated_root=tmp_path / "skills"
        )

        sandbox = LocalDockerSandbox(workspace_root=tmp_path / "wsp")
        try:
            session_id = f"owner_t10_pdf:conv_{uuid.uuid4().hex[:8]}"
            await sandbox.create_session(
                session_id, limits=ResourceLimits(), network=NetworkPolicy()
            )

            chat_script: list[ChatResponse] = [
                # Step 0: research preamble (web_research mock).
                _resp(
                    tool_calls=[
                        ToolCall(
                            name="web_search",
                            args={"query": "tenant protection norway"},
                            call_id="ws1",
                        )
                    ]
                ),
                # Step 1: activate pdf_generation.
                _resp(
                    tool_calls=[
                        ToolCall(
                            name="use_skill",
                            args={"skill_name": "pdf_generation"},
                            call_id="us1",
                        )
                    ]
                ),
                # Step 2: code_execution with the reportlab snippet.
                _resp(
                    tool_calls=[
                        ToolCall(
                            name="code_execution",
                            args={"code": _PDF_CODE},
                            call_id="ce1",
                        )
                    ]
                ),
                # Step 3: final.
                _resp("[FINAL] Wrote /workspace/out/report.pdf"),
            ]
            backend = _ScriptedBackend(chat_script=chat_script)
            loop = _build_agentic_loop(
                backend,
                sandbox=sandbox,
                session_id=session_id,
                scanned=scanned,
                persona=_persona(),
                extra_tools=[_mock_web_search, _mock_web_fetch],
            )

            run = await loop.run("research X then produce a PDF report")

            assert run.status is RunStatus.COMPLETED
            # The final code_execution step is the penultimate step (the last
            # step is FINAL); assert that step surfaced the produced .pdf in
            # its ToolResult.data.
            code_step = next(
                s
                for s in run.steps
                if s.type is StepType.TOOL_CALL
                and any(c.name == "code_execution" for c in s.tool_calls)
            )
            ce_result = next(
                r
                for r, c in zip(code_step.results, code_step.tool_calls, strict=True)
                if c.name == "code_execution"
            )
            assert ce_result.is_error is False, f"code_execution failed: {ce_result.content!r}"
            assert ce_result.data is not None
            produced = ce_result.data.get("produced_files") or []
            assert any(p.get("path", "").endswith("report.pdf") for p in produced), (
                f"expected report.pdf in produced_files; got {produced}"
            )

            # And the host-side workspace really has it (the sandbox is
            # source-of-truth — ExecutionResult.produced_files is built from
            # the snapshot-then-diff over host_out per D-12-10).
            host_out = sandbox._session_workspaces[session_id][1]
            pdf_files = list(host_out.glob("**/*.pdf"))
            assert len(pdf_files) == 1
            assert pdf_files[0].name == "report.pdf"
        finally:
            await sandbox.aclose()


class TestErrorRecovery:
    """Criterion #9 — broken python-docx code recovers via tool-error feedback."""

    @pytest.mark.asyncio
    async def test_bad_python_docx_code_recovers_via_tool_error_recovery(
        self, tmp_path: Path
    ) -> None:
        # Criterion #9. Mirrors test_loop_agentic.py:161's
        # test_tool_failure_is_fed_back_and_model_recovers — but the failure
        # is a python-docx AttributeError traceback from inside the sandbox,
        # not a flaky-tool 503.
        skip_reason = _docker_skip_reason()
        if skip_reason:
            pytest.skip(skip_reason)
        scanned = _scanned_skills(["docx_generation"], isolated_root=tmp_path / "skills")

        sandbox = LocalDockerSandbox(workspace_root=tmp_path / "wsp")
        try:
            session_id = f"owner_t10_recovery:conv_{uuid.uuid4().hex[:8]}"
            await sandbox.create_session(
                session_id, limits=ResourceLimits(), network=NetworkPolicy()
            )

            chat_script: list[ChatResponse] = [
                # Step 0: activate docx_generation.
                _resp(
                    tool_calls=[
                        ToolCall(
                            name="use_skill",
                            args={"skill_name": "docx_generation"},
                            call_id="us1",
                        )
                    ]
                ),
                # Step 1: BROKEN code → sandbox surfaces AttributeError.
                _resp(
                    tool_calls=[
                        ToolCall(
                            name="code_execution",
                            args={"code": _BROKEN_DOCX_CODE},
                            call_id="ce1",
                        )
                    ]
                ),
                # Step 2: model sees the error, writes FIXED code.
                _resp(
                    tool_calls=[
                        ToolCall(
                            name="code_execution",
                            args={"code": _FIXED_DOCX_CODE},
                            call_id="ce2",
                        )
                    ]
                ),
                # Step 3: final.
                _resp("[FINAL] Recovered and wrote /workspace/out/memo.docx"),
            ]
            backend = _ScriptedBackend(chat_script=chat_script)
            loop = _build_agentic_loop(
                backend,
                sandbox=sandbox,
                session_id=session_id,
                scanned=scanned,
                persona=_persona(),
            )

            run = await loop.run("draft a memo as a docx")

            # The turn did NOT crash — the loop fed the broken-code error
            # back as a ToolResult(is_error=True, ...); the next round
            # produced the fixed snippet.
            assert run.status is RunStatus.COMPLETED, (
                f"expected COMPLETED after recovery; got {run.status}: {run.error!r}"
            )

            # Find the two code_execution steps + their results.
            code_steps = [
                s
                for s in run.steps
                if s.type is StepType.TOOL_CALL
                and any(c.name == "code_execution" for c in s.tool_calls)
            ]
            assert len(code_steps) == 2, (
                f"expected 2 code_execution steps (broken + fixed); got {len(code_steps)}"
            )

            broken_result = code_steps[0].results[0]
            assert broken_result.is_error is True
            # The traceback content contains the AttributeError message OR
            # the outcome=error marker from the result-formatter; either
            # signals to the model that the call failed.
            content_lower = broken_result.content.lower()
            assert "attributeerror" in content_lower or "outcome=error" in content_lower, (
                f"expected an error signal in {broken_result.content!r}"
            )

            fixed_result = code_steps[1].results[0]
            assert fixed_result.is_error is False, (
                f"recovery snippet should succeed: {fixed_result.content!r}"
            )
            assert fixed_result.data is not None
            produced = fixed_result.data.get("produced_files") or []
            assert any(p.get("path", "").endswith("memo.docx") for p in produced), (
                f"expected memo.docx in produced_files; got {produced}"
            )

            host_out = sandbox._session_workspaces[session_id][1]
            assert list(host_out.glob("**/memo.docx")), (
                "memo.docx should exist on the host after recovery"
            )
        finally:
            await sandbox.aclose()


class TestTenantWorkspaceIsolation:
    """Criterion #11 — tenant-isolated session_ids land in separate workspaces.

    Inherits Spec 12 D-12-1 (tenant-prefixed session_id) + D-12-9 (two-mount
    workspace) + D-12-15 (filesystem isolation precise definition). Each
    ``LocalDockerSandbox`` session owns its own per-session ``host_out``
    bind-mount directory, so owner_B literally cannot reach owner_A's file
    by path (different host directory, different container, different mount).
    """

    @pytest.mark.asyncio
    async def test_docx_workspace_isolated_per_tenant(self, tmp_path: Path) -> None:
        skip_reason = _docker_skip_reason()
        if skip_reason:
            pytest.skip(skip_reason)

        owner_a = "owner_A_t10"
        owner_b = "owner_B_t10"
        conv_id = f"conv_{uuid.uuid4().hex[:8]}"  # SAME conversation_id on both
        session_a = f"{owner_a}:{conv_id}"
        session_b = f"{owner_b}:{conv_id}"

        sandbox = LocalDockerSandbox(workspace_root=tmp_path / "wsp")
        try:
            # 1. Two sessions, distinct tenant-prefixed session_ids.
            await sandbox.create_session(
                session_a, limits=ResourceLimits(), network=NetworkPolicy()
            )
            await sandbox.create_session(
                session_b, limits=ResourceLimits(), network=NetworkPolicy()
            )

            # 2. owner_A writes a docx in session_A.
            result_a = await sandbox.execute(
                _TENANT_A_WRITE_CODE,
                session_id=session_a,
                timeout_s=30.0,
            )
            assert result_a.outcome == "ok", (
                f"tenant_A write failed: stdout={result_a.stdout!r} stderr={result_a.stderr!r}"
            )
            assert any(f.path.endswith("tenant_a_memo.docx") for f in result_a.produced_files), (
                f"tenant_a_memo.docx missing from {result_a.produced_files}"
            )

            # 3. Per Spec 12 D-12-9 / D-12-15: the two sessions' host-side
            # workspace directories must be distinct paths. This is the
            # structural isolation proof — independent of any in-container
            # check (different host dir → different bind-mount → no shared
            # filesystem state).
            host_in_a, host_out_a = sandbox._session_workspaces[session_a]
            host_in_b, host_out_b = sandbox._session_workspaces[session_b]
            assert host_out_a != host_out_b, (
                "owner_A and owner_B must have distinct host workspace dirs"
            )
            assert host_in_a != host_in_b
            # And session_a's docx is reachable on the host (sanity).
            assert (host_out_a / "tenant_a_memo.docx").is_file()

            # 4. owner_B's session probes /workspace/out: assert tenant_a's
            # file is NOT visible. The probe also enumerates the directory
            # so an accidental shared-mount regression would surface in
            # the listing output (debuggable failure mode).
            result_b = await sandbox.execute(
                _TENANT_B_READ_PROBE_CODE,
                session_id=session_b,
                timeout_s=30.0,
            )
            assert result_b.outcome == "ok", f"tenant_B probe failed: stderr={result_b.stderr!r}"
            assert "exists: False" in result_b.stdout, (
                f"owner_B should NOT see owner_A's docx; got stdout={result_b.stdout!r}"
            )
            assert "tenant_a_memo.docx" not in result_b.stdout, (
                f"owner_A's filename leaked into owner_B's listing: stdout={result_b.stdout!r}"
            )
        finally:
            await sandbox.aclose()


# ---------------------------------------------------------------------------
# Hosted-backend scoping marker (D-16-X-5 / A2). Documented sentinel so a
# future test author who adds a hosted parametrisation sees the explicit
# scope decision — the test below is a no-op skip whose body explains why.
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_hosted_backend_produced_files_not_yet_wired() -> None:
    """Sentinel: HostedSandbox does NOT populate ``produced_files`` at v0.1.

    Per state.md A2: ``packages/api/src/persona_api/sandbox/hosted.py``'s
    ``_run_and_marshal`` constructs ``ExecutionResult`` with the default
    empty ``produced_files`` tuple. Criterion #8 / #9 are scoped to the
    local backend for T10. When Spec 17's D-17-X-bytes-persistence lands
    the bytes-from-sandbox bridge for hosted, this skip is removed and
    a parametrisation across both backends is added.
    """
    pytest.skip(
        "HostedSandbox.produced_files is unwired at v0.1 (state.md A2 / "
        "D-16-X-5). T10 is scoped to LocalDockerSandbox only."
    )

"""Spec 16 T09 — allow-list gating (#4) + composition (#7) integration test.

Exercises the criterion-#4 (allow-list gating: a persona without
``docx_generation`` declared cannot invoke it) and criterion-#7 (content/
format composition: ``document_drafting`` → in-context prose → ``use_skill
(docx_generation)`` → ``code_execution`` produces a real ``.docx`` from
the prose-as-Python-string, with NO ``.md`` intermediate) acceptance
criteria using a real :class:`Toolbox` + the real
:func:`persona.skills.make_use_skill_tool` factory + the real
:class:`LocalDockerSandbox`.

Per T01 audit A1 (state.md): the api package has no shared ``_fakes.py``;
the established pattern is module-private ``_ScriptedBackend`` per test
file (three precedents — test_runtime_factory.py, test_authoring.py,
test_api_authoring_service.py). We copy the
:class:`persona_runtime.tests._fakes.ScriptedBackend` shape inline rather
than cross-import (brittle; crosses package boundaries).

Per T01 audit A2 / D-16-X-5: the produced-files contract surfaces at the
sandbox-boundary level (``LocalDockerSandbox._discover_produced_files``
populates ``ExecutionResult.produced_files`` BEFORE the per-execution
``shutil.rmtree(host_out)`` cleanup). This test asserts at the sandbox
boundary via the tool result's ``data["produced_files"]`` (the textual
mirror of ``ExecutionResult.produced_files``) — NOT at the persona-
workspace persistence layer (cross-session persistence is deferred to v0.2
per D-16-X-5).

Markers: ``integration`` (Docker daemon + image) AND ``docker`` (declared
in pyproject.toml ``[tool.pytest.ini_options].markers``) so CI hosts
without Docker can opt out via ``-m "not docker"`` cleanly without losing
the wider ``integration`` selection.
"""

# ruff: noqa: ANN401, ARG002, D102 — test doubles with intentionally loose sigs.

from __future__ import annotations

import json
from pathlib import Path
from typing import TYPE_CHECKING, Any

import pytest
from persona.backends.types import ChatResponse, TokenUsage
from persona.sandbox import make_code_execution_tool
from persona.sandbox.local_docker import (
    DEFAULT_IMAGE,
    LocalDockerSandbox,
    is_docker_available,
)
from persona.sandbox.result import NetworkPolicy, ResourceLimits
from persona.schema.persona import Persona, PersonaIdentity
from persona.schema.skills import SkillSpec
from persona.schema.tools import ToolCall, ToolResult
from persona.skills import (
    SkillInjector,
    SkillScanner,
    count_tokens,
    make_use_skill_tool,
)
from persona.tools import Toolbox
from persona_runtime.agentic.loop import AgenticLoop
from persona_runtime.agentic.run import RunStatus
from persona_runtime.agentic.step import StepType
from persona_runtime.prompt import PromptBuilder
from persona_runtime.router import Router
from persona_runtime.tier import TierConfig, TierRegistry

if TYPE_CHECKING:
    from persona.backends.types import ToolSpec
    from persona.schema.conversation import ConversationMessage

pytestmark = [pytest.mark.integration, pytest.mark.docker]


# The bundled built-in skills live alongside persona-core's source. The
# scanner resolves user paths first, then built-ins; we point at built-ins
# directly so the test composes real SKILL.md packs.
_BUILTIN_ROOT = (
    Path(__file__).parent.parent.parent.parent / "core" / "src" / "persona" / "skills" / "builtin"
).resolve()

# Minimal in-sandbox python-docx snippet — keeps the integration test fast
# while exercising the real library and proving the produced-files contract
# round-trips a real .docx. The body mirrors what a real persona would
# produce after activating ``docx_generation``: prose-as-Python-string
# embedded directly (NOT a .md file read back in), per D-16-3.
_DRAFTED_PROSE = (
    "Norwegian tenancy protections shield tenants from unilateral rent "
    "increases. Notice periods extended to six months in 2024."
)

# D-12-X-venv-path-ordering (Spec 16 T09/T10 fix): the LocalDockerSandbox
# container PATH now prepends ``/opt/venv/bin`` so the persona-sandbox
# image's venv-installed ``python-docx`` is reachable natively. The
# ``sys.path.insert`` workaround that earlier T09 drafts carried is no
# longer needed — SKILL.md packs (and this composition test) can import
# ``from docx import Document`` directly.
_DOCX_CODE_TEMPLATE = """\
from docx import Document
doc = Document()
doc.add_heading('Tenant Protection Memo', level=1)
doc.add_paragraph({prose!r})
out_path = '/workspace/out/tenant_memo.docx'
doc.save(out_path)
print(f'wrote: {{out_path}}')
"""


def _image_present(tag: str) -> bool:
    """True if the sandbox image is available on the local Docker daemon."""
    import docker
    from docker.errors import DockerException, ImageNotFound

    try:
        client = docker.from_env()
        try:
            client.images.get(tag)
        except ImageNotFound:
            return False
        finally:
            client.close()
    except DockerException:
        return False
    return True


# ----- module-private scripted backend (per-file pattern, T01 audit A1) -----


class _ScriptedBackend:
    """Replays a scripted sequence of :class:`ChatResponse` objects.

    Copied shape from ``packages/runtime/tests/_fakes.py:50`` per T01 audit
    A1 (the api package convention is module-private per-file, not a
    shared fixture). Only the ``chat()`` non-streaming surface the
    :class:`AgenticLoop` drives is implemented; ``chat_stream`` raises
    because the AgenticLoop never calls it.
    """

    def __init__(
        self,
        chat_script: list[ChatResponse],
        *,
        provider_name: str = "anthropic",
        model_name: str = "claude-sonnet-4-6",
    ) -> None:
        self._chat_script = list(chat_script)
        self._chat_index = 0
        self._provider_name = provider_name
        self._model_name = model_name
        self.chat_calls = 0

    @property
    def provider_name(self) -> str:
        return self._provider_name

    @property
    def model_name(self) -> str:
        return self._model_name

    @property
    def supports_native_tools(self) -> bool:
        return False

    @property
    def supports_vision(self) -> bool:
        return False

    async def chat(
        self,
        messages: list[ConversationMessage],
        *,
        tools: list[ToolSpec] | None = None,
        temperature: float = 0.0,
        max_tokens: int = 4096,
        stop: list[str] | None = None,
    ) -> Any:
        self.chat_calls += 1
        if self._chat_index < len(self._chat_script):
            response = self._chat_script[self._chat_index]
            self._chat_index += 1
            return response
        # Defensive empty final — the loop should have terminated by here.
        return ChatResponse(
            content="[FINAL] ",
            tool_calls=[],
            usage=TokenUsage(prompt_tokens=1, completion_tokens=0, total_tokens=1),
            model=self._model_name,
            provider=self._provider_name,
            latency_ms=0.0,
        )

    async def chat_stream(
        self,
        messages: list[ConversationMessage],
        *,
        tools: list[ToolSpec] | None = None,
        temperature: float = 0.0,
        max_tokens: int = 4096,
        stop: list[str] | None = None,
    ) -> Any:
        # The AgenticLoop never calls chat_stream (it uses chat() only).
        raise NotImplementedError("AgenticLoop drives chat(), not chat_stream()")


# ----- builders -------------------------------------------------------------


def _persona(*, skills: list[str], tools: list[str]) -> Persona:
    """Build a minimal :class:`Persona` for the integration test.

    Identity carries the required fields; no ``visual_style`` (the
    docx_generation skill body short-circuits cleanly when absent).
    """
    return Persona(
        persona_id="tenancy_assistant",
        identity=PersonaIdentity(
            name="Astrid",
            role="Norwegian tenancy assistant",
            background="Helps tenants understand husleieloven in plain language.",
            language_default="nb",
            constraints=["Do not give binding legal advice."],
        ),
        skills=skills,
        tools=tools,
    )


def _resp(content: str = "", *, tool_calls: list[ToolCall] | None = None) -> ChatResponse:
    return ChatResponse(
        content=content,
        tool_calls=tool_calls or [],
        usage=TokenUsage(prompt_tokens=10, completion_tokens=5, total_tokens=15),
        model="claude-sonnet-4-6",
        provider="anthropic",
        latency_ms=1.0,
    )


def _empty_stores() -> dict[str, Any]:
    """Minimal in-memory stores; the test asserts on tool results, not memory."""
    from typing import Any as _Any

    class _MemStore:  # noqa: D401 — local test double, no docstring needed
        def write(self, *_a: _Any, **_kw: _Any) -> None: ...
        def query(self, *_a: _Any, **_kw: _Any) -> list[_Any]:
            return []

        def get_all(self, *_a: _Any, **_kw: _Any) -> list[_Any]:
            return []

        def delete(self, *_a: _Any, **_kw: _Any) -> None: ...
        def remove_documents(self, *_a: _Any, **_kw: _Any) -> None: ...
        def history(self, *_a: _Any, **_kw: _Any) -> list[_Any]:
            return []

        def rollback(self, *_a: _Any, **_kw: _Any) -> None: ...

    return {
        "identity": _MemStore(),
        "self_facts": _MemStore(),
        "worldview": _MemStore(),
        "episodic": _MemStore(),
    }


def _build_loop(
    *,
    persona: Persona,
    backend: _ScriptedBackend,
    toolbox: Toolbox,
    scanned_skills: list[SkillSpec],
    max_steps: int = 20,
) -> AgenticLoop:
    """Wire an :class:`AgenticLoop` over the supplied collaborators.

    The scripted backend stands in for every tier; the registry's lazy-
    instantiation cache is primed directly (same idiom as
    ``packages/runtime/tests/unit/test_loop_agentic.py:_make_loop``).
    """
    # ruff: noqa: SLF001 — pinning the registry cache to the scripted backend
    # mirrors the runtime test idiom; the alternative is a real BackendConfig
    # plus a no-op load_backend monkeypatch, which is brittler.
    from persona.backends import BackendConfig

    dummy_cfg = BackendConfig(provider="anthropic", model="m", api_key=None)  # type: ignore[arg-type]
    registry = TierRegistry(
        {
            "frontier": TierConfig(name="frontier", backend_config=dummy_cfg),
            "mid": TierConfig(name="mid", backend_config=dummy_cfg),
            "small": TierConfig(name="small", backend_config=dummy_cfg),
        }
    )
    registry._cache = {  # type: ignore[assignment]
        "frontier": backend,  # type: ignore[dict-item]
        "mid": backend,  # type: ignore[dict-item]
        "small": backend,  # type: ignore[dict-item]
    }
    return AgenticLoop(
        persona=persona,
        stores=_empty_stores(),  # type: ignore[arg-type]
        toolbox=toolbox,
        skill_injector=SkillInjector(),
        scanned_skills=scanned_skills,
        prompt_builder=PromptBuilder(),
        router=Router(),
        tier_registry=registry,
        max_steps=max_steps,
    )


def _scan(skill_names: list[str]) -> list[SkillSpec]:
    """Scan the named built-in skills, asserting all resolved."""
    scanner = SkillScanner([_BUILTIN_ROOT])
    specs = scanner.scan(skill_names)
    assert len(specs) == len(skill_names), (
        f"scanner returned {len(specs)} specs for {skill_names!r}; "
        "expected all to resolve from the built-in path"
    )
    return specs


# ----- tests ----------------------------------------------------------------


class TestAllowListGating:
    """Criterion #4 — a persona without ``docx_generation`` cannot invoke it."""

    @pytest.mark.asyncio
    async def test_persona_without_docx_generation_cannot_invoke(self) -> None:
        """A persona declaring only ``web_research`` calling
        ``use_skill('docx_generation')`` receives a structured error from the
        ``make_use_skill_tool`` factory — the tool is registered (the
        runtime composes it for any persona that has skills) but the
        requested skill name is not in the closure-captured ``available``
        set (Spec 04 ``use_skill_tool.py`` ``Unknown skill: ...; available:
        ...`` shape).

        Per state.md A1 + the task brief: the assertion targets the exact
        phrasing the shipped ``make_use_skill_tool`` produces — so a future
        refactor of the error message gets caught by this test rather than
        silently changing the persona-facing contract.
        """
        scanned = _scan(["web_research"])
        use_skill = make_use_skill_tool(scanned)
        toolbox = Toolbox([use_skill], allow_list=["use_skill"])

        call = ToolCall(
            name="use_skill",
            args={"skill_name": "docx_generation"},
            call_id="c1",
        )
        result: ToolResult = await toolbox.dispatch(call)

        assert result.is_error is True, (
            "use_skill should return is_error=True for a skill that isn't in the "
            "persona's scanned skills (the persona declared only web_research)"
        )
        # Shape verbatim from persona.skills.use_skill_tool.make_use_skill_tool
        # — the f-string template is "Unknown skill: {skill_name}; available:
        # {', '.join(sorted(available)) or '(none)'}".
        assert "Unknown skill: docx_generation" in result.content, (
            f"expected 'Unknown skill: docx_generation' in content; got: {result.content!r}"
        )
        assert "available:" in result.content, (
            f"expected 'available:' phrasing in content; got: {result.content!r}"
        )
        assert "web_research" in result.content, (
            f"expected the available list to include web_research; got: {result.content!r}"
        )


class TestCompositionEndToEnd:
    """Criterion #7 — content/format composition end-to-end.

    Drives an :class:`AgenticLoop` through a four-step scripted scenario:

    a. Tool call ``use_skill('document_drafting')`` → activated.
    b. Text turn produces drafted prose IN model context (assert NO file
       written by drafting — D-16-3 verification: drafting writes to
       context, NOT a .md file the next skill reads back in).
    c. Tool call ``use_skill('docx_generation')`` → activated.
    d. Tool call ``code_execution`` with python-docx code embedding the
       drafted prose as a Python string literal (the canonical
       composition pattern D-16-3 prescribes).
    e. Assert ``ExecutionResult.produced_files`` (surfaced via the tool
       result's ``data['produced_files']``) contains exactly one ``.docx``
       — sandbox-boundary level per T01 audit A2 scoping (the
       persona-workspace persistence chain is out of scope; D-16-X-5).
    """

    @pytest.mark.asyncio
    async def test_document_drafting_then_docx_generation_composes_end_to_end(
        self,
        tmp_path: Path,
    ) -> None:
        if not is_docker_available():
            pytest.skip(
                "Docker daemon unreachable; LocalDockerSandbox integration "
                "test requires a running Docker daemon."
            )
        if not _image_present(DEFAULT_IMAGE):
            pytest.skip(
                f"Sandbox image {DEFAULT_IMAGE!r} not built. Run the T06 "
                "build (or `docker pull`) before exercising the LocalDocker "
                "composition path."
            )

        persona = _persona(
            skills=["document_drafting", "docx_generation"],
            tools=["use_skill", "code_execution"],
        )
        scanned = _scan(["document_drafting", "docx_generation"])

        # Real LocalDockerSandbox + real make_code_execution_tool — the
        # composition test is meaningful only when the actual sandbox runs
        # the python-docx code.
        #
        # Use SESSION MODE (D-12-1 scaled scope): create_session spawns a
        # long-lived container; D-12-X-venv-path-ordering ensures both
        # session-mode ``docker exec`` and one-shot ``docker run`` resolve
        # ``python`` from ``/opt/venv/bin`` so the image's installed
        # ``docx``/``pptx``/``openpyxl``/``reportlab`` are importable
        # natively without any sys.path workaround.
        sandbox = LocalDockerSandbox(workspace_root=tmp_path / "sbx")
        session_id = f"{persona.persona_id}:conv_t09"
        try:
            await sandbox.create_session(
                session_id, limits=ResourceLimits(), network=NetworkPolicy()
            )
            code_exec = make_code_execution_tool(
                sandbox,
                persona_id=persona.persona_id,
                session_id_provider=lambda: session_id,
            )
            use_skill = make_use_skill_tool(scanned)
            toolbox = Toolbox(
                [use_skill, code_exec],
                allow_list=["use_skill", "code_execution"],
            )

            docx_code = _DOCX_CODE_TEMPLATE.format(prose=_DRAFTED_PROSE)
            chat_script = [
                # (a) Activate document_drafting.
                _resp(
                    tool_calls=[
                        ToolCall(
                            name="use_skill",
                            args={"skill_name": "document_drafting"},
                            call_id="c1",
                        )
                    ]
                ),
                # (b) Drafting produces prose IN context (no file write —
                # D-16-3: the bridge is the model's own context, NOT a .md
                # file chain). Note: NO tool call, plain text — the loop
                # treats this as a REASONING step and proceeds.
                _resp(
                    f"Draft: {_DRAFTED_PROSE} "
                    "Now I will activate docx_generation to produce the .docx."
                ),
                # (c) Activate docx_generation.
                _resp(
                    tool_calls=[
                        ToolCall(
                            name="use_skill",
                            args={"skill_name": "docx_generation"},
                            call_id="c2",
                        )
                    ]
                ),
                # (d) code_execution embedding the drafted prose as a
                # python string literal (composition WITHOUT a .md
                # intermediate).
                _resp(
                    tool_calls=[
                        ToolCall(
                            name="code_execution",
                            args={"code": docx_code},
                            call_id="c3",
                        )
                    ]
                ),
                # (e) Final.
                _resp("[FINAL] The .docx has been produced and is ready."),
            ]
            backend = _ScriptedBackend(chat_script)
            loop = _build_loop(
                persona=persona,
                backend=backend,
                toolbox=toolbox,
                scanned_skills=scanned,
            )

            run = await loop.run("draft a tenant-protection memo and produce a .docx")

            # Run terminated cleanly.
            assert run.status is RunStatus.COMPLETED, (
                f"run did not complete: status={run.status} error={run.error!r}"
            )

            # Step shape: 4 tool/reasoning steps + 1 final step.
            step_types = [s.type for s in run.steps]
            assert step_types == [
                StepType.TOOL_CALL,  # (a) use_skill(document_drafting)
                StepType.REASONING,  # (b) drafted prose in context
                StepType.TOOL_CALL,  # (c) use_skill(docx_generation)
                StepType.TOOL_CALL,  # (d) code_execution(python-docx)
                StepType.FINAL,  # (e) [FINAL]
            ], f"unexpected step type sequence: {step_types}"

            # D-16-3 verification: the drafting step produced prose in the
            # model's content, NOT a file. The drafting tool call's result
            # carries data={'skill_name': 'document_drafting'} (the
            # use_skill activation envelope) — the actual prose lives in
            # the SUBSEQUENT REASONING step's content.
            drafting_call = run.steps[0]
            assert drafting_call.tool_calls[0].name == "use_skill"
            assert drafting_call.tool_calls[0].args == {"skill_name": "document_drafting"}
            assert drafting_call.results[0].is_error is False
            assert drafting_call.results[0].data == {"skill_name": "document_drafting"}

            drafting_reasoning = run.steps[1]
            assert drafting_reasoning.type is StepType.REASONING
            assert _DRAFTED_PROSE in (drafting_reasoning.content or ""), (
                "drafted prose must appear in the REASONING step's content "
                "(D-16-3: bridge = model context, NOT a .md file)"
            )

            # D-16-3 verification continued: the .docx-producing
            # code_execution step embeds the drafted prose as a Python
            # string literal — NOT as a file read. Inspect the call args.
            docx_gen_call = run.steps[2]
            assert docx_gen_call.tool_calls[0].args == {"skill_name": "docx_generation"}

            code_call_step = run.steps[3]
            assert code_call_step.tool_calls[0].name == "code_execution"
            code_arg = code_call_step.tool_calls[0].args["code"]
            assert _DRAFTED_PROSE in code_arg, (
                "the prose drafted in step (b) must be embedded in the "
                "python-docx code as a string literal (D-16-3 composition)"
            )
            assert ".read_text(" not in code_arg, (
                "D-16-3: composition must NOT chain via Path(...).read_text(); "
                "the bridge is the model's own context (a Python string literal)"
            )
            assert ".md" not in code_arg, (
                "D-16-3: composition must NOT reference a .md intermediate; "
                "the bridge is the model's own context (a Python string literal)"
            )

            # Sandbox-boundary assertion (T01 audit A2 scoping): the
            # code_execution tool result mirrors
            # ExecutionResult.produced_files into ToolResult.data
            # ['produced_files'] BEFORE the LocalDockerSandbox per-execution
            # host_out cleanup. We assert at this boundary; the
            # persona-workspace persistence chain is out of scope per
            # D-16-X-5 (cross-session reachability is a v0.2 follow-up).
            code_result = code_call_step.results[0]
            assert code_result.is_error is False, (
                f"code_execution should have succeeded: {code_result.content!r}"
            )
            produced = (code_result.data or {}).get("produced_files") or []
            assert isinstance(produced, list)
            docx_files = [pf for pf in produced if str(pf.get("path", "")).endswith(".docx")]
            assert len(docx_files) == 1, (
                f"expected exactly one .docx in produced_files; got: {json.dumps(produced)}"
            )
            assert docx_files[0]["size_bytes"] > 0, (
                f".docx produced_files entry has zero size: {docx_files[0]!r}"
            )

            # Composition criterion #7: docx_generation's body is the
            # injected skill content the loop appended to context AFTER
            # the use_skill tool result (the loop's _maybe_inject_skill
            # bridge). Sanity-check that the scanned spec carries the
            # expected token budget — a regression guard symmetric to
            # builtin-skills test_each_under_token_budget.
            docx_spec = next(s for s in scanned if s.name == "docx_generation")
            assert count_tokens(docx_spec.content) <= 2000
        finally:
            await sandbox.aclose()
